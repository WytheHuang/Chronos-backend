import io
from datetime import UTC
from datetime import datetime
from typing import Any
from uuid import UUID

import docx
from httpx import delete
import pytz
from asgiref.sync import sync_to_async
from core import exceptions as core_exceptions
from core import schemas as core_schemas
from core.apis import BaseEditApiController
from core.models import User
from django.conf import settings
from django.contrib.auth.models import AnonymousUser
from django.core.handlers.asgi import ASGIRequest
from django.http import FileResponse
from django_celery_results.models import TaskResult
from langchain.chains.question_answering import load_qa_chain
from ninja import Form
from ninja.files import UploadedFile
from ninja_extra import api_controller
from ninja_extra import route
from ninja_extra.permissions import IsAuthenticated
from ninja_jwt.authentication import AsyncJWTAuth
from qdrant_client.http import models as rest
from django.db.utils import IntegrityError

from . import models
from . import schemas
from . import tasks
from . import utils


@api_controller(
    prefix_or_class="/file-upload-logs",
    auth=AsyncJWTAuth(),
    tags=["File Upload Logs"],
    permissions=[IsAuthenticated],
)
class FileUploadLogModelEditController(BaseEditApiController):
    """FileUploadLogModelEditController."""

    Model = models.FileUploadLogModel

    @route.get(
        "",
        response={
            200: list[schemas.FileUploadLogModelSchema],
            401: core_schemas.Http401UnauthorizedSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def get_all(
        self,
        request: ASGIRequest,  # noqa: ARG002
    ) -> list[schemas.FileUploadLogModelSchema]:
        """Get all file upload logs.

        Args:
            request (ASGIRequest): django request.

        Returns:
            list[schemas.FileUploadLogModelSchema]: file upload logs.
        """
        # logs = await sync_to_async(models.FileUploadLogModel.objects.all)()
        logs = await sync_to_async(models.FileUploadLogModel.objects.all().order_by)("-created_at")

        res = []
        async for log in logs:
            try:
                task = await TaskResult.objects.aget(task_id=log.task_id)
                embeddings_status = task.status
            except TaskResult.DoesNotExist:
                embeddings_status = str(None)

            try:
                user = await User.objects.aget(id=log.upload_by_id)  # type: ignore
                upload_by = user.email
            except User.DoesNotExist:
                upload_by = str(None)

            try:
                data_category = await models.CategoryModel.objects.aget(id=log.category_id)  # type: ignore
            except models.CategoryModel.DoesNotExist:
                data_category = None

            data = {
                "id": log.id,
                "s3_key": log.s3_key,
                "created_at": log.created_at.astimezone(tz=pytz.timezone("Asia/Taipei")).isoformat(),
                "embeddings_status": embeddings_status,
                "category": data_category.name if data_category else "None",  # type: ignore
                "file_name": log.file_name,
                "upload_by": upload_by,
            }

            await sync_to_async(res.append)(data)

        return res

    @route.get(
        "{id}",
        response={
            200: schemas.FileUploadLogModelSchema,
            401: core_schemas.Http401UnauthorizedSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def get(
        self,
        request: ASGIRequest,
        id: UUID,  # noqa: A002
    ) -> schemas.FileUploadLogModelSchema:
        """Get file upload log by id.

        Args:
            request (ASGIRequest): django request.
            id (UUID): file upload log id.

        Returns:
            schemas.FileUploadLogModelSchema: file upload log.
        """
        return await sync_to_async(super().get)(request, id)


@api_controller(
    prefix_or_class="/categories",
    auth=AsyncJWTAuth(),
    tags=["Categories"],
    permissions=[IsAuthenticated],
)
class CategoryModelEditController(BaseEditApiController):
    """CategoryModelEditController."""

    Model = models.CategoryModel

    @route.get(
        "",
        response={
            200: list[schemas.CategoryModelSchema],
            401: core_schemas.Http401UnauthorizedSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def get_all(
        self,
        request: ASGIRequest,  # noqa: ARG002
    ) -> list[dict[str, Any]]:
        """Get all categories.

        Args:
            request (ASGIRequest): django request.

        Returns:
            list[schemas.CategoryModelSchema]: categories.
        """
        model = await sync_to_async(models.CategoryModel.objects.all().values)()

        return await sync_to_async(list)(model)

    @route.post(
        "",
        response={
            200: schemas.CategoryModelSchema,
            401: core_schemas.Http401UnauthorizedSchema,
            403: core_schemas.Http403ForbiddenSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def create(
        self,
        request: ASGIRequest,
        body: schemas.CreateCategoryRequestSchema,
    ) -> schemas.CategoryModelSchema:
        """Create a category.

        Args:
            request (ASGIRequest): django request.
            name (str): category name.

        Returns:
            schemas.CategoryModelSchema: category.
        """
        if isinstance(request.user, AnonymousUser):
            raise core_exceptions.Http401UnauthorizedException
        if not request.user.is_staff:  # type: ignore
            raise core_exceptions.Http403ForbiddenException

        try:
            response = await sync_to_async(super().create)(request, body)
        except IntegrityError as err:
            raise core_exceptions.Http400BadRequestException(detail=str(err))

        return response

    @route.delete(
        "",
        response={
            200: core_schemas.BaseResponseSchema,
            401: core_schemas.Http401UnauthorizedSchema,
            403: core_schemas.Http403ForbiddenSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def delete(self, request: ASGIRequest, id: UUID) -> dict[str, Any]:
        """Delete a category.

        Args:
            request (ASGIRequest): django request.
            id (UUID): category id.

        Returns:
            core_schemas.BaseResponseSchema: category.
        """
        if isinstance(request.user, AnonymousUser):
            raise core_exceptions.Http401UnauthorizedException
        if not request.user.is_staff:  # type: ignore
            raise core_exceptions.Http403ForbiddenException

        try:
            category = await sync_to_async(models.CategoryModel.objects.get)(id=id)
        except models.CategoryModel.DoesNotExist as err:
            raise core_exceptions.Http400BadRequestException(detail="category_id not found") from err

        files_in_category = await sync_to_async(models.FileUploadLogModel.objects.filter)(category=category)
        delete_status = []
        async for file in files_in_category:
            delete_status.append(await sync_to_async(utils.delete_file)(file.s3_key))
            await sync_to_async(file.delete)(request.user)

        return {
            "file_delete": "acknowledged" not in delete_status,
            **await sync_to_async(super().delete)(request, id),
        }

    @route.delete(
        "/v2",
        response={
            200: core_schemas.BaseResponseSchema,
            401: core_schemas.Http401UnauthorizedSchema,
            403: core_schemas.Http403ForbiddenSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def delete_v2(self, request: ASGIRequest, id: UUID) -> dict[str, Any]:
        """Delete a category.

        Args:
            request (ASGIRequest): django request.
            id (UUID): category id.

        Returns:
            core_schemas.BaseResponseSchema: category.
        """
        if isinstance(request.user, AnonymousUser):
            raise core_exceptions.Http401UnauthorizedException
        if not request.user.is_staff:  # type: ignore
            raise core_exceptions.Http403ForbiddenException

        try:
            category = await sync_to_async(models.CategoryModel.objects.get)(id=id)
        except models.CategoryModel.DoesNotExist as err:
            raise core_exceptions.Http400BadRequestException(detail="category_id not found") from err

        files_in_category = await sync_to_async(models.FileUploadLogModel.objects.filter)(category=category)
        delete_status = await sync_to_async(utils.delete_category)(category.name)
        await sync_to_async(files_in_category.update)(
            is_delete=True,
            deleted_by_user=request.user,
            deleted_at=datetime.now(tz=pytz.timezone("Asia/Taipei")),
        )

        return {
            "file_delete": delete_status == "completed",
            **await sync_to_async(super().delete)(request, id),
        }


@api_controller(
    prefix_or_class="/qa",
    auth=AsyncJWTAuth(),
    tags=["QA"],
    permissions=[IsAuthenticated],
)
class QAApiController:
    """Controller for handling API requests related to question answering, document generation, and content download.

    This class manages multiple endpoints, each serving a different purpose:
    - '/ask' for processing user queries for question answering.
    - '/document' for generating documents using an AI model.
    - '/download' for downloading content based on specific criteria.
    """

    @route.post(
        "/ask",
        response={
            200: schemas.AskDocumentResponseSchema,
            401: core_schemas.Http401UnauthorizedSchema,
        },
    )
    async def ask(
        self,
        request: ASGIRequest,  # noqa: ARG002
        body: schemas.AskDocumentRequestSchema,
    ):
        """Asynchronously processes POST requests to the '/ask' endpoint for question answering.

        This method handles incoming questions using the AskDocumentRequestSchema and provides
        responses as per the AskDocumentResponseSchema. It manages successful and unauthorized requests
        with 200 and 401 status codes, respectively.

        Args:
            request (ASGIRequest): The ASGI request object representing the HTTP request.
            body (AskDocumentRequestSchema): The schema containing the user's question details.

        Returns:
            Response object: Either AskDocumentResponseSchema for successful queries or
            Http401UnauthorizedSchema for unauthorized access.
        """

        def _get_filters(
            filter_categories: list[str],
            filter_docs: list[str],
        ) -> rest.Filter:
            """Constructs a filter object for use in REST queries.

            This method takes lists of categories and documents, along with optional metadata keys,
            to create a filter object. This filter can then be used to refine search or query operations
            in REST-based services.

            Args:
                filter_categories (list[str]): A list of category names to include in the filter.
                filter_docs (list[str]): A list of document identifiers to include in the filter.
                metadata_payload_key (str, optional): The key in the payload where metadata is stored.
                    Defaults to "metadata".
                metadata_root_key (str, optional): The root key under which the metadata is categorized.
                    Defaults to "source".

            Returns:
                rest.Filter: A filter object configured based on the provided categories, documents,
                and metadata keys.
            """
            should_conditions = []
            if filter_categories:
                should_conditions.append(
                    rest.FieldCondition(
                        key=f"{settings.METADATA_PAYLOAD_KEY}.{settings.METADATA_ROOT_KEY}.分類名稱",
                        match=rest.MatchAny(any=filter_categories),
                    ),
                )
            if filter_docs:
                should_conditions.append(
                    rest.FieldCondition(
                        key=f"{settings.METADATA_PAYLOAD_KEY}.{settings.METADATA_ROOT_KEY}.文件名稱",
                        match=rest.MatchAny(any=filter_docs),
                    ),
                )
            return rest.Filter(should=should_conditions)

        def _main_qa(
            query: str,
            filter_categories: list[str],
            filter_docs: list[str],
            k: int = 3,
        ) -> dict[str, Any] | None:
            """Processes a question-answering query and returns relevant information.

            This method takes a user's query and applies optional filters based on categories
            and documents to refine the search. It returns a dictionary with the original question,
            the generated answer, and references to documents that are relevant to the answer.
            If no relevant answer is found, or if the processing fails, the method may return None.

            Args:
                query (str): The user's question or query string.
                filter_categories (list[str]): Categories to filter the answers by.
                filter_docs (list[str]): Document identifiers to refine the search.
                k (int, optional): The maximum number of answers to consider. Defaults to 5.

            Returns:
                dict[str, Any] | None: A dictionary with keys 'Question', 'Answer', and 'References',
                where 'Question' is the original query, 'Answer' is the generated response,
                and 'References' are the related documents. Returns None if no suitable answer is found.
            """
            if utils.Qdrant_VectorStore:
                filter_condition = _get_filters(
                    filter_categories=filter_categories,
                    filter_docs=filter_docs,
                )
                chain = load_qa_chain(llm=utils.LLM_MODEL, chain_type="stuff", prompt=utils.PROMPT)
                docsearch_docs = utils.Qdrant_VectorStore.similarity_search(
                    query=str(query),
                    k=k,
                    filter=filter_condition,
                )
                # result = chain({"input_documents": docsearch_docs, "question": query}, return_only_outputs=False)
                return chain.invoke({"input_documents": docsearch_docs, "question": query}, return_only_outputs=False)

            else:  # noqa: RET505
                return None

        result = _main_qa(
            query=body.question,
            filter_categories=body.filter_categories,
            filter_docs=body.filter_docs,
        )

        if result is not None:
            reference_docs = {}
            for index, document in enumerate(result["input_documents"]):
                reference_docs[index] = {
                    "page_content": document.page_content,
                    "metadata": document.metadata,
                }
            return {"Question": body.question, "Answer": result["output_text"], "References": reference_docs}
        else:  # noqa: RET505
            return {"Question": body.question, "Answer": "Error", "References": {}}

    @route.post(
        "/document",
        response={
            200: schemas.GenerativeAIDocumentResponseSchema,
            400: core_schemas.Http400BadRequestSchema,
            401: core_schemas.Http401UnauthorizedSchema,
        },
    )
    async def document(  # noqa: C901, PLR0912
        self,
        request: ASGIRequest,
        body: schemas.GenerativeAIDocumentRequestSchema,
        model: str = settings.OPENAI_GPT_35_MODEL_NAME,
    ):
        """Asynchronously processes requests for document generation at the '/document' endpoint.

        This endpoint uses an AI model, specified in the request or defaulting to the application's
        settings, to generate documents based on the provided schema. It returns responses based on
        the GenerativeAIDocumentResponseSchema, and handles bad requests and unauthorized access
        with 400 and 401 status codes, respectively.

        Args:
            request (ASGIRequest): The incoming ASGI request object.
            body (GenerativeAIDocumentRequestSchema): Schema detailing the document generation requirements.
            model (str, optional): The AI model used for document generation, defaulting to the application's setting.

        Returns:
            Response object: Either GenerativeAIDocumentResponseSchema for successful generation,
            Http400BadRequestSchema for bad requests, or Http401UnauthorizedSchema for unauthorized access.
        """
        if isinstance(request.user, AnonymousUser):
            raise core_exceptions.Http401UnauthorizedException

        question = body.question
        documents = body.documents

        if len(documents) > 3:  # noqa: PLR2004
            raise core_exceptions.Http400BadRequestException(detail="Too many documents")

        combined_text = ""
        if len(documents) == 0:
            combined_text = "There are no documents. So the task for you is to follow the instructions from the user."
        else:
            for index, doc_name in enumerate(documents):
                try:
                    f_queryset = await sync_to_async(models.FileUploadLogModel.objects.filter)(file_name=doc_name)
                    f = await sync_to_async(f_queryset.first)()
                    if f:
                        combined_text += f"\n - \nThis is the content of file number {index+1}, file name is '{doc_name}'. Below is the content of this file.\n\n"
                        docx_content = await sync_to_async(utils.s3.get_object)(
                            Bucket=settings.AWS_S3_BUCKET_NAME,
                            Key=f.s3_key,
                        )
                        doc_file = docx_content["Body"].read()

                        if ".pdf" in doc_name:
                            doc_file = utils.pdf_to_docx(doc_file)
                        if isinstance(doc_file, bytes):
                            doc_file = utils.bytes_to_docx(doc_file)

                        for paragraph in doc_file.paragraphs:
                            combined_text += paragraph.text
                    else:
                        combined_text += ""

                except models.FileUploadLogModel.DoesNotExist:
                    raise core_exceptions.Http400BadRequestException(  # noqa: B904, TRY200
                        detail=f"Document {doc_name} not found",
                    )

        try:
            messages = [
                {
                    "role": "system",
                    "content": """
                        You are a document assistant.
                        Your task is to modify specific content in a document or add new content in specified areas according to user requests.
                        Your responses should include only the modified or newly added content of the document.
                        Please ensure to use Traditional Chinese when creating new documents in the .docx format.""",
                },
                {
                    "role": "user",
                    "content": f"""
                    你的任務是，要依照使用者的需求針對特定文件的內容進行修改，並回傳更改過的段落內容。
                    -
                    Here are the instructions: {question}. Here are the file contents: {combined_text}.

                    -
                    你必須遵守以下事項：
                    1. 根據用戶指示修改特定文件。
                    2. 保持文件的原始內容不變，僅更改並回傳用戶指定的部分，不需要整份文件都回傳給客戶。
                    3. 以.docx格式提供修改後的文件，包含必要的文件元素，如段落、標題等。
                    4. 如果用戶的請求超出修改文件的範圍，則回答「不清楚該請求，請再輸入一次。」。
                    5. 禁止使用簡體中文或是英文等其他語言回答，請以繁體中文（正體中文）進行回答。
                    """,
                },
            ]

            response = utils.OPENAI_client.chat.completions.create(
                model=model,
                messages=messages,  # type: ignore
                temperature=settings.OPENAI_TEMPERATURE,
            )
            response_content = response.choices[0].message.content if response else ""
            finish_reason = response.choices[0].finish_reason if response else ""

            while finish_reason == "length":
                messages = [
                    *messages,
                    {
                        "role": "assistant",
                        "content": response_content[-300:] if response_content != "" else "",  # type: ignore
                    },
                    {
                        "role": "user",
                        "content": """
                        你尚未完成請繼續剛才的任務，繼續完成文件的創建和修改。
                        """,
                    },
                ]
                response = utils.OPENAI_client.chat.completions.create(
                    model=model,
                    messages=messages,  # type: ignore
                    temperature=settings.OPENAI_TEMPERATURE,
                )

                response_content += response.choices[0].message.content if response else ""  # type: ignore
                finish_reason = response.choices[0].finish_reason if response else ""

            file_stream = io.BytesIO()
            doc = docx.Document()
            doc.add_paragraph(response_content)
            doc.save(file_stream)
            file_stream.seek(0)

            doc_names = "_".join([name.split(".")[0] for name in documents])
            concatenated_names = (
                f"{datetime.now(tz=pytz.timezone('Asia/Taipei')).strftime('%Y%m%d_%H%M%S')}_[{doc_names}]_生成結果.docx"
            )

            file_info = {
                "category": settings.FILE_STANDARD_CATEGORY_NAME,
                "filename": concatenated_names,
                "upload_time": datetime.now(tz=pytz.timezone("Asia/Taipei")).isoformat(),
            }

            s3_key = await sync_to_async(utils.get_s3_key)(file_info)
            utils.s3.upload_fileobj(file_stream, settings.AWS_S3_BUCKET_NAME, s3_key)

            return {  # noqa: TRY300
                "message": response_content,
                "s3_key": s3_key,
            }

        except Exception:  # noqa: BLE001
            return {
                "message": "Some Error Occurred",
                "s3_key": "Error",
            }

    @route.post(
        "/download",
        response={
            200: Any,
            401: core_schemas.Http401UnauthorizedSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def download_by_string(
        self,
        request: ASGIRequest,
        body: schemas.DownloadByStringRequestSchema,
    ):
        """Asynchronously handles POST requests to the '/download' endpoint for content download.

        This method processes requests for downloading content based on criteria specified
        in the DownloadByStringRequestSchema. It returns the requested content for successful
        requests, and handles unauthorized requests and not found errors with 401 and 404
        status codes, respectively.

        Args:
            request (ASGIRequest): The incoming ASGI request object.
            body (DownloadByStringRequestSchema): The schema containing the criteria for downloading content.

        Returns:
            On success, the requested content in the specified format.
            On unauthorized access, an Http401UnauthorizedSchema response.
            On content not found, an Http404NotFoundSchema response.
        """
        if isinstance(request.user, AnonymousUser):
            raise core_exceptions.Http401UnauthorizedException

        file_stream = io.BytesIO()
        doc = docx.Document()
        doc.add_paragraph(body.string)
        doc.save(file_stream)
        file_stream.seek(0)

        return FileResponse(file_stream, as_attachment=True, filename="generated.docx")


@api_controller(
    prefix_or_class="/file",
    auth=AsyncJWTAuth(),
    tags=["file"],
    permissions=[IsAuthenticated],
)
class FileApiController:
    """Controller for handling file-related API requests including upload, download, and checking the status of embedding tasks.

    This class manages endpoints for various file operations:
    - '/upload' for uploading files and initiating embedding tasks.
    - '/download' for downloading files based on their storage key.
    - '/embedding-status' for retrieving the status of an embedding task.
    """

    @route.post(
        "/upload",
        response={
            200: schemas.FileUploadResponseSchema,
            400: core_schemas.Http400BadRequestSchema,
            401: core_schemas.Http401UnauthorizedSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def upload(self, request: ASGIRequest, file: UploadedFile, category_id: Form[UUID]):
        """Asynchronously handles file uploads and initiates an embedding task.

        Receives a file and an optional category, processes the file upload, and starts an
        embedding task. Returns the task ID for further tracking.

        Args:
            request (ASGIRequest): The incoming ASGI request object.
            file (UploadedFile): The file to be uploaded.
            category (Form[UUID]): The category_id of the file.

        Returns:
            dict: A dictionary containing the embedding task ID.
        """
        if isinstance(request.user, AnonymousUser):
            raise core_exceptions.Http401UnauthorizedException

        past_files = await sync_to_async(models.FileUploadLogModel.objects.filter)(file_name=file.name)
        if await sync_to_async(len)(past_files) > 0:
            raise core_exceptions.Http400BadRequestException("filename already exists")

        if file.name.split(".")[-1] not in ["pdf", "docx", "doc"]:
            raise core_exceptions.Http400BadRequestException("file type not supported")

        try:
            category = await sync_to_async(models.CategoryModel.objects.get)(id=category_id)
        except models.CategoryModel.DoesNotExist as err:
            raise core_exceptions.Http400BadRequestException(detail="category_id not found") from err

        file_info = {
            "category": category.name,
            "uploaded_by": str(request.user.id),  # type: ignore
            "filename": file.name,
            "upload_time": datetime.now(UTC).astimezone(tz=pytz.timezone("Asia/Taipei")).isoformat(),
        }
        s3_key = await sync_to_async(utils.get_s3_key)(file_info)

        await sync_to_async(utils.s3.upload_fileobj)(file.file, settings.AWS_S3_BUCKET_NAME, s3_key)

        task_id = tasks.embding_task.apply_async(kwargs={"s3_key": s3_key})

        log = models.FileUploadLogModel(
            file_name=file.name,
            s3_key=s3_key,
            upload_by=request.user,
            task_id=str(task_id),
            category=category,
        )
        await sync_to_async(log.create)(request.user)

        return {"embedding_task_id": str(task_id)}

    @route.delete(
        "",
        response={
            200: core_schemas.BaseResponseSchema,
            401: core_schemas.Http401UnauthorizedSchema,
            403: core_schemas.Http403ForbiddenSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def delete(self, request: ASGIRequest, s3_key: str):
        if isinstance(request.user, AnonymousUser):
            raise core_exceptions.Http401UnauthorizedException
        if not (request.user.is_staff or utils.get_file_info(s3_key)["uploaded_by"] == str(request.user.id)):  # type: ignore
            raise core_exceptions.Http403ForbiddenException

        try:
            model = await models.FileUploadLogModel.objects.aget(s3_key=s3_key)
        except models.FileUploadLogModel.DoesNotExist as err:
            raise core_exceptions.Http400BadRequestException(detail="s3_key not found") from err

        delete_status = await sync_to_async(utils.delete_file)(s3_key)
        await sync_to_async(model.delete)(request.user)

        return {"msg": "success", "delete_status": delete_status}

    @route.get(
        "/download",
        response={
            200: Any,
            401: core_schemas.Http401UnauthorizedSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def download_by_key(
        self,
        request: ASGIRequest,
        s3_key: str,
    ):
        """Asynchronously handles file download requests based on an S3 key.

        Retrieves and returns a file stored in S3, identified by the provided S3 key. The file
        is returned as an attachment.

        Args:
            request (ASGIRequest): The incoming ASGI request object.
            s3_key (str): The S3 key used to locate and retrieve the file.

        Returns:
            FileResponse: A response containing the file as an attachment.
        """
        if isinstance(request.user, AnonymousUser):
            raise core_exceptions.Http401UnauthorizedException

        s3_obj = await sync_to_async(utils.s3.get_object)(
            Bucket=settings.AWS_S3_BUCKET_NAME,
            Key=s3_key,
        )
        file_info = await sync_to_async(utils.get_file_info)(s3_key)

        return FileResponse(s3_obj["Body"], as_attachment=True, filename=file_info["filename"])

    @route.get(
        "/embedding-status",
        response={
            200: schemas.EmbeddingStatusResponseSchema,
            401: core_schemas.Http401UnauthorizedSchema,
            404: core_schemas.Http404NotFoundSchema,
        },
    )
    async def embedding_status(self, request: ASGIRequest, task_id: UUID):
        """Asynchronously checks the status of an embedding task.

        Receives an embedding task ID and returns the current status of the task.

        Args:
            request (ASGIRequest): The incoming ASGI request object.
            task_id (UUID): The unique identifier of the embedding task.

        Returns:
            dict: A dictionary containing the status of the embedding task. basicaly is "SUCCESS", "FAILURE" or "STARTED"
        """
        if isinstance(request.user, AnonymousUser):
            raise core_exceptions.Http401UnauthorizedException

        try:
            task_rusult = await sync_to_async(TaskResult.objects.get)(task_id=task_id)
        except TaskResult.DoesNotExist as err:
            raise core_exceptions.Http404NotFoundException(detail="task_id not found") from err

        return {"status": task_rusult.status}
