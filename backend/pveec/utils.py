import io
from ast import literal_eval
from typing import Any
from typing import Literal

import boto3
import pdfplumber
import requests
from cryptography.fernet import Fernet
from django.conf import settings
from docx import Document
from langchain.prompts import PromptTemplate
from langchain_community.vectorstores import Qdrant
from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings
from openai import OpenAI
from qdrant_client import QdrantClient

from qdrant_client.http import models as rest

from . import utils


__encoder = Fernet(str(settings.FILE_INFO_SECRET_KEY).encode("utf-8"))

s3 = boto3.client(
    "s3",
    aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
    aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
)
OPENAI_client = OpenAI(api_key=settings.OPENAI_API_KEY)

# Embedding Object
OPENAI_EMBEDDING = OpenAIEmbeddings(
    # openai_api_key=settings.OPENAI_API_KEY,
)
LLM_MODEL = ChatOpenAI(
    # openai_api_key=settings.OPENAI_API_KEY,
    model=settings.OPENAI_GPT_4_MODEL_NAME,
    temperature=settings.OPENAI_TEMPERATURE,
    max_tokens=settings.OPENAI_MAX_TOKENS,
)  # type: ignore


# Qdrant
def url_check(url: str):
    """Checks if the given URL is accessible and returns a 200 OK status.

    Args:
        url (str): The URL to be checked.

    Returns:
        bool: True if the URL returns a 200 OK status, False otherwise.
    """
    response = requests.get(url)  # noqa: S113
    return response.status_code == 200  # noqa: PLR2004


def list_collections(url: str) -> list[str | None]:
    """Retrieves a list of collection names from a Qdrant service.

    Args:
        url (str): The base URL of the Qdrant service.

    Returns:
        List[Optional[str]]: A list of collection names. Returns a list containing
        None if the request fails.
    """
    endpoint = f"{url}/collections"
    response = requests.get(endpoint)  # noqa: S113

    if response.status_code == 200:  # noqa: PLR2004
        collections = response.json().get("result", {}).get("collections", [])
        return [collection["name"] for collection in collections]
    else:  # noqa: RET505
        return [None]


def _get_exsisting_qdrant_client(
    qdrant_url: str,
) -> QdrantClient:
    """
    Retrieves an existing Qdrant client from the specified URL.

    Args:
        qdrant_url (str): The base URL of the Qdrant service.

    Returns:
        Qdrant: An instance of Qdrant representing the retrieved embedding.

    Raises:
        ValueError: If the Qdrant URL is not accessible.
    """
    if url_check(qdrant_url):
        return QdrantClient(url=qdrant_url)

    raise ValueError("Qdrant URL is not accessible.")


Qdrant_client = _get_exsisting_qdrant_client(
    settings.QDRANT_URL,
)


def _get_exsisting_embedding_qdrant(
    qdrant_url: str,
    collection_name: str,
) -> Qdrant:
    """Retrieves an existing embedding from the specified collection in Qdrant.

    Args:
        qdrant_url (str): The base URL of the Qdrant service.
        collection_name (str): The name of the collection to retrieve the embedding from.

    Returns:
        Qdrant: An instance of Qdrant representing the retrieved embedding.
    """
    if url_check(qdrant_url) and collection_name in list_collections(url=qdrant_url):
        return Qdrant(Qdrant_client, collection_name, OPENAI_EMBEDDING)
    return False


Qdrant_VectorStore = _get_exsisting_embedding_qdrant(
    settings.QDRANT_URL,
    settings.QDRANT_COLLECTION,
)

# QA
prompt_template = """
您是「台灣傑美工程顧問股份有限公司」的專業問答助理，專注於回答有關土壤和地下水污染調查、整治以及相關環境工程領域與公司內行政事宜如請假、工作流程等等問題。

請利用公司的背景資料和您的專業知識來回答問題。如果您不確定答案，請直接說明不知道，而不要提供不確定的資訊。請務必使用繁體中文進行回答。

{context}

Question: {question}

請使用繁體中文進行回答:"""
PROMPT = PromptTemplate(
    template=prompt_template,
    input_variables=["context", "question"],
)


# Type Convertor
def pdf_to_docx(pdf_bytes: Any):
    """Converts PDF byte content to a DOCX document.

    Args:
        pdf_bytes (Any): Byte content of the PDF to be converted.

    Returns:
        Any: The converted DOCX document.
    """
    pdf_file = io.BytesIO(pdf_bytes)
    doc = Document()
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                cleaned_text = "".join(char for char in text if char.isprintable() and char != "\x00")
                doc.add_paragraph(cleaned_text)

    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    pdf_file.close()
    return doc


def bytes_to_docx(data: Any) -> Any:
    """Converts byte content to a DOCX document.

    Args:
        data (Any): Byte content to be converted into a DOCX document.

    Returns:
        Any: A DOCX document created from the byte content.
    """
    data = io.BytesIO(data)
    return Document(docx=data)


# Functions
def __get_file_info_encode(file_info: dict) -> str:
    """Encodes and encrypts file information into a string.

    Args:
        file_info (dict): The file information to be encoded and encrypted.

    Returns:
        str: The encrypted string representation of the file information.
    """
    return __encoder.encrypt(str(file_info).encode("utf-8")).decode("utf-8")


def __get_file_info_decode(file_info_hash: str) -> dict:
    """Decrypts and decodes file information from a string.

    Args:
        file_info_hash (str): The encrypted string representation of the file information.

    Returns:
        dict: The decrypted and decoded file information.
    """
    return literal_eval(__encoder.decrypt(file_info_hash.encode("utf-8")).decode("utf-8"))


def get_s3_key(file_info: dict) -> str:
    """Generates an S3 key for a file based on its information.

    Args:
        file_info (dict): The information about the file.

    Returns:
        str: A string representing the S3 key for the file.
    """
    return f"{__get_file_info_encode(file_info)}.{file_info['filename'].split('.')[-1]}"


def get_file_info(s3_key: str) -> dict:
    """Retrieves file information from an S3 key.

    This function decodes the file information encoded in the S3 key. It splits the S3 key
    to extract the encoded file information and then decodes it to retrieve the original
    file information.

    Args:
        s3_key (str): The S3 key of the file, which contains encoded file information.

    Returns:
        dict: The decoded file information.
    """
    return __get_file_info_decode(s3_key.split(".")[0])


def delete_file(s3_key: str) -> Literal["acknowledged", "completed"]:
    """Deletes a file from the S3 bucket.

    Args:
        s3_key (str): The S3 key of the file to be deleted.
    """
    delete_status = utils.Qdrant_client.delete(
        collection_name=settings.QDRANT_COLLECTION,
        points_selector=rest.FilterSelector(
            filter=rest.Filter(
                must=[
                    rest.FieldCondition(
                        key=f"{settings.METADATA_PAYLOAD_KEY}.{settings.METADATA_ROOT_KEY}.文件名稱",
                        match=rest.MatchValue(value=utils.get_file_info(s3_key)["filename"]),
                    ),
                ],
            )
        ),
    )

    return delete_status.status.value


def delete_category(category_name: str) -> Literal["acknowledged", "completed"]:
    delete_status = utils.Qdrant_client.delete(
        collection_name=settings.QDRANT_COLLECTION,
        points_selector=rest.FilterSelector(
            filter=rest.Filter(
                must=[
                    rest.FieldCondition(
                        key=f"{settings.METADATA_PAYLOAD_KEY}.{settings.METADATA_ROOT_KEY}.分類名稱",
                        match=rest.MatchValue(value=category_name),
                    ),
                ],
            )
        ),
    )

    return delete_status.status.value
